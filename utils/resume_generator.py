# utils/resume_generator.py
import os
import uuid
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
# Import XML elements in a way that's compatible with different python-docx versions
try:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn

# Import docx2pdf conditionally to avoid errors if it's not available
try:
    import docx2pdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

def add_horizontal_line(paragraph):
    """Add a horizontal line to a paragraph using a simpler method"""
    paragraph.add_run('_' * 80)

def generate_resume_file(data):
    # Create a unique filename
    unique_id = str(uuid.uuid4())
    docx_path = f"resumes/resume_{unique_id}.docx"
    
    # Check if an existing resume was uploaded
    uploaded_resume_path = data.get("uploaded_resume_path")
    
    # Ensure the resumes directory exists
    os.makedirs("resumes", exist_ok=True)
    
    # Create a new Document
    doc = Document()
    
    # Set document margins (narrower margins for more content space)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Set up styles
    styles = doc.styles
    
    # Name style
    name_style = styles.add_style('ResumeName', WD_STYLE_TYPE.PARAGRAPH)
    name_style.font.name = 'Calibri'
    name_style.font.size = Pt(24)  # Larger name
    name_style.font.bold = True
    name_style.font.color.rgb = RGBColor(0, 59, 113)  # Professional blue
    name_style.paragraph_format.space_after = Pt(0)
    
    # Job Title style
    job_title_style = styles.add_style('ResumeJobTitle', WD_STYLE_TYPE.PARAGRAPH)
    job_title_style.font.name = 'Calibri'
    job_title_style.font.size = Pt(14)
    job_title_style.font.italic = True
    job_title_style.font.color.rgb = RGBColor(68, 68, 68)  # Dark gray
    job_title_style.paragraph_format.space_after = Pt(6)
    
    # Contact style
    contact_style = styles.add_style('ResumeContact', WD_STYLE_TYPE.PARAGRAPH)
    contact_style.font.name = 'Calibri'
    contact_style.font.size = Pt(10)
    contact_style.font.color.rgb = RGBColor(68, 68, 68)  # Dark gray
    contact_style.paragraph_format.space_after = Pt(12)
    
    # Section Heading style
    heading_style = styles.add_style('ResumeHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Calibri'
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 59, 113)  # Professional blue
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)
    
    # Subheading style
    subheading_style = styles.add_style('ResumeSubheading', WD_STYLE_TYPE.PARAGRAPH)
    subheading_style.font.name = 'Calibri'
    subheading_style.font.size = Pt(12)
    subheading_style.font.bold = True
    subheading_style.paragraph_format.space_before = Pt(6)
    subheading_style.paragraph_format.space_after = Pt(0)
    
    # Date style
    date_style = styles.add_style('ResumeDate', WD_STYLE_TYPE.PARAGRAPH)
    date_style.font.name = 'Calibri'
    date_style.font.size = Pt(10)
    date_style.font.italic = True
    date_style.font.color.rgb = RGBColor(102, 102, 102)  # Medium gray
    date_style.paragraph_format.space_after = Pt(3)
    
    # Normal text style
    normal_style = styles.add_style('ResumeNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)
    
    # Bullet style
    bullet_style = styles.add_style('ResumeBullet', WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.font.name = 'Calibri'
    bullet_style.font.size = Pt(11)
    bullet_style.paragraph_format.left_indent = Inches(0.25)
    bullet_style.paragraph_format.space_after = Pt(3)
    
    # Add name as title
    name = doc.add_paragraph(data.get("full_name", "Unnamed"), style='ResumeName')
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add field of work and experience level
    job_title = doc.add_paragraph(style='ResumeJobTitle')
    job_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    job_title.add_run(f"{data.get('field_of_work', '')} - {data.get('experience_level', '')} ({data.get('years_of_experience', '')} years)")
    
    # Add contact information
    contact_info = doc.add_paragraph(style='ResumeContact')
    contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_parts = []
    if data.get('email'):
        contact_parts.append(data.get('email'))
    if data.get('phone'):
        contact_parts.append(data.get('phone'))
    if data.get('location'):
        contact_parts.append(data.get('location'))
    if data.get('linkedin'):
        contact_parts.append(f"LinkedIn: {data.get('linkedin')}")
    
    contact_info.add_run(' | '.join(contact_parts))
    
    # Add other social media links if available
    social_links = []
    if data.get('github'):
        social_links.append(f"GitHub: {data.get('github')}")
    if data.get('website'):
        social_links.append(f"Website: {data.get('website')}")
    if data.get('twitter'):
        social_links.append(f"Twitter: {data.get('twitter')}")
    
    if social_links:
        social_info = doc.add_paragraph(style='ResumeContact')
        social_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        social_info.add_run(' | '.join(social_links))
    
    # Add a line separator
    separator = doc.add_paragraph()
    add_horizontal_line(separator)
    
    # Professional Summary
    doc.add_paragraph("PROFESSIONAL SUMMARY", style='ResumeHeading')
    summary_para = doc.add_paragraph(data.get("summary", ""), style='ResumeNormal')
    
    # Skills
    doc.add_paragraph("SKILLS", style='ResumeHeading')
    skills_text = data.get("skills", "")
    
    # Format skills as a clean list
    skills_list = [skill.strip() for skill in skills_text.split(',') if skill.strip()]
    
    # Add skills as bullet points
    for skill in skills_list:
        skill_para = doc.add_paragraph(style='ResumeBullet')
        skill_para.add_run(f"• {skill}")
    
    # Work Experience
    doc.add_paragraph("WORK EXPERIENCE", style='ResumeHeading')
    experience_text = data.get("experience", "")
    for exp in experience_text.split("|"):
        if exp.strip():
            parts = exp.split(",", 2)
            if len(parts) >= 2:
                company = parts[0].strip()
                title = parts[1].strip()
                
                # Company and position
                exp_para = doc.add_paragraph(style='ResumeSubheading')
                exp_para.add_run(f"{company}").bold = True
                
                # Position as a separate line
                position_para = doc.add_paragraph(style='ResumeDate')
                position_para.add_run(title)
                
                # Description with bullet points
                if len(parts) > 2:
                    desc = parts[2].strip()
                    # Split description into bullet points if it contains semicolons
                    if ';' in desc:
                        for bullet in desc.split(';'):
                            if bullet.strip():
                                bullet_para = doc.add_paragraph(style='ResumeBullet')
                                bullet_para.add_run(f"• {bullet.strip()}")
                    else:
                        # Otherwise add as a normal paragraph
                        doc.add_paragraph(desc, style='ResumeNormal')
            else:
                doc.add_paragraph(exp.strip(), style='ResumeNormal')
    
    # Education
    doc.add_paragraph("EDUCATION", style='ResumeHeading')
    education_text = data.get("education", "")
    for edu in education_text.split("|"):
        if edu.strip():
            parts = edu.split(",")
            if len(parts) >= 2:
                degree = parts[0].strip()
                institution = parts[1].strip()
                
                # Institution name
                edu_para = doc.add_paragraph(style='ResumeSubheading')
                edu_para.add_run(institution).bold = True
                
                # Degree info
                degree_para = doc.add_paragraph(style='ResumeNormal')
                degree_para.add_run(degree)
                
                # Year if available
                if len(parts) > 2:
                    year = parts[2].strip()
                    year_para = doc.add_paragraph(style='ResumeDate')
                    year_para.add_run(f"Graduation: {year}")
            else:
                doc.add_paragraph(edu.strip(), style='ResumeNormal')
    
    # Projects (Optional)
    projects = data.get("projects")
    if projects and projects.strip():
        doc.add_paragraph("PROJECTS", style='ResumeHeading')
        
        for project in projects.split("|"):
            if project.strip():
                parts = project.split(",", 1)
                if len(parts) >= 1:
                    project_name = parts[0].strip()
                    proj_para = doc.add_paragraph(style='ResumeSubheading')
                    proj_para.add_run(project_name).bold = True
                    
                    if len(parts) > 1:
                        desc = parts[1].strip()
                        # Split description into bullet points if it contains semicolons
                        if ';' in desc:
                            for bullet in desc.split(';'):
                                if bullet.strip():
                                    bullet_para = doc.add_paragraph(style='ResumeBullet')
                                    bullet_para.add_run(f"• {bullet.strip()}")
                        else:
                            doc.add_paragraph(desc, style='ResumeNormal')
    
    # Online Presence (Optional)
    website = data.get("website")
    blog = data.get("blog")
    youtube = data.get("youtube")
    github = data.get("github")
    
    if website or blog or youtube or github:
        doc.add_paragraph("ONLINE PRESENCE", style='ResumeHeading')
        
        online_links = []
        if website:
            online_links.append(f"Personal Website: {website}")
        if blog:
            online_links.append(f"Blog: {blog}")
        if youtube:
            online_links.append(f"YouTube Channel: {youtube}")
        if github:
            online_links.append(f"GitHub: {github}")
            
        for link in online_links:
            link_para = doc.add_paragraph(style='ResumeBullet')
            link_para.add_run(f"• {link}")
    
    # Certifications (Optional)
    certifications = data.get("certifications")
    if certifications and certifications.strip():
        doc.add_paragraph("CERTIFICATIONS", style='ResumeHeading')
        
        # Format certifications as bullet points
        cert_list = [cert.strip() for cert in certifications.split(',') if cert.strip()]
        for cert in cert_list:
            cert_para = doc.add_paragraph(style='ResumeBullet')
            cert_para.add_run(f"• {cert}")
    
    # Languages (Optional)
    languages = data.get("languages")
    if languages and languages.strip():
        doc.add_paragraph("LANGUAGES", style='ResumeHeading')
        
        # Format languages as bullet points
        lang_list = [lang.strip() for lang in languages.split(',') if lang.strip()]
        for lang in lang_list:
            lang_para = doc.add_paragraph(style='ResumeBullet')
            lang_para.add_run(f"• {lang}")
    
    # Add a simple footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_para.add_run("Resume generated by Resume Kraft")
    footer_run.font.size = Pt(9)
    footer_run.font.name = 'Calibri'
    
    # Save the document
    doc.save(docx_path)
    
    # Convert to PDF if requested - FEATURE TEMPORARILY DISABLED
    if data.get("output_format") == "pdf":
        print("PDF conversion temporarily disabled. Returning DOCX instead.")
        return docx_path
    
    return docx_path
